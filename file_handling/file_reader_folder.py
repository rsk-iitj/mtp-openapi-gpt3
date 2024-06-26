import os
import base64
from docx import Document
import re
from io import BytesIO
import pandas as pd
import pdfplumber



def extract_text_from_file(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    try:
        if ext == '.docx':
            doc = Document(file_path)
            return '\n'.join([paragraph.text for paragraph in doc.paragraphs if paragraph.text])
        elif ext in ['.txt', '.md']:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        elif ext == '.pdf':
            with pdfplumber.open(file_path) as pdf:
                pages = [page.extract_text() for page in pdf.pages]
                return '\n'.join(pages) if pages else None
        elif ext in ['.xls', '.xlsx']:
            # Read the Excel file and concatenate all cells into a single string
            df = pd.read_excel(file_path)
            return df.to_string(header=True, index=False)
        elif ext == '.doc':
            # Handling .doc files requires additional dependencies and setup
            return handle_doc_file(file_path)
    except Exception as e:
        print(f"Failed to read {file_path}: {str(e)}")
    return None

def handle_doc_file(file_path):
    # This is a placeholder for handling .doc files; you might need to adjust based on your environment
    # textract approach (requires installation of antiword or similar on non-Windows systems)
    import textract
    return textract.process(file_path).decode('utf-8')

def extract_texts_from_folder(directory):
    supported_formats = ['.txt', '.md', '.pdf', '.xls', '.xlsx', '.doc', '.docx']
    texts = []
    file_names = []
    for root, dirs, files in os.walk(directory):
        for file in files:
            if any(file.lower().endswith(ext) for ext in supported_formats):
                file_path = os.path.join(root, file)
                file_text = extract_text_from_file(file_path)
                if file_text:
                    texts.append(file_text)
                    file_names.append(file)
    return texts, file_names




def save_test_plan(full_test_plan, application_name, model_used):
    # Define the directory structure
    base_directory = "output/test-plan"
    application_directory = os.path.join(base_directory, application_name.replace(" ", "_"))

    # Create the directory if it does not exist
    os.makedirs(application_directory, exist_ok=True)

    # Initialize the Word document
    doc = Document()
    for section, content in full_test_plan.items():
        # Add main section heading from the dictionary key
        doc.add_heading(section, level=1)

        # Split content by lines for fine-grained processing
        lines = content.split('\n')
        for line in lines:
            # Skip empty lines
            if not line.strip():
                continue
            # Handling Markdown-style headings within the section content
            if line.startswith('## '):
                doc.add_heading(line.replace('## ', ''), level=2)
            elif line.startswith('### '):
                doc.add_heading(line.replace('### ', ''), level=3)
            elif line.startswith('#### '):
                doc.add_heading(line.replace('#### ', ''), level=4)
            else:
                # Process the paragraph and handle bold formatting
                p = doc.add_paragraph()
                for part in re.split(r'(\*\*[^*]+\*\*)', line):  # Split and keep the bold parts
                    if part.startswith('**') and part.endswith('**'):
                        part = part[2:-2]  # Remove the asterisks
                        p.add_run(part).bold = True
                    else:
                        p.add_run(part)

    # Save the document
    filename = os.path.join(application_directory, f"{application_name.replace(' ', '_')}_Test_Plan_{model_used}.docx")
    doc.save(filename)
    print(f"Document saved to {filename}")
    return doc


def download_link(doc, filename, text):
    # Generate download link for the document
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)  # Reset buffer position to the start of the stream
    b64 = base64.b64encode(buffer.getvalue()).decode()
    href = f'<a href="data:application/octet-stream;base64,{b64}" download="{filename}">{text}</a>'
    return href